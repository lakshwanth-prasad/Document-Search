#Imports
import docx
import re
import nltk
import fitz
from django.shortcuts import render,redirect
from django.core.files.storage import FileSystemStorage
from .models import Document, PositionalIndex
from PyPDF2 import PdfReader
from nltk.tokenize import word_tokenize
from nltk.stem.porter import PorterStemmer
from django.utils.html import escape, mark_safe

nltk.download('punkt')
#iew for handling the upload form and saving uploaded documents.
def home(request):
    if request.method == 'POST':
        files = request.FILES.getlist('documents')
        for file in files:
            fs = FileSystemStorage()
            filename = fs.save(file.name, file)
            file_path = fs.path(filename)
            text = extract_text(file_path, file.name)
            document = Document(file=fs.url(filename), content=text)
            document.save()
            process_text_and_index(text, document.id)
        return render(request, 'upload_complete.html', {'documents': Document.objects.all()})
    return render(request, 'upload.html')

#Similar to 'home' view, processes document uploads and indexes them.  
def upload_and_index(request):
    if request.method == 'POST':
        files = request.FILES.getlist('documents')
        for file in files:
            fs = FileSystemStorage()
            filename = fs.save(file.name, file)
            file_path = fs.path(filename)
            text = extract_text(file_path, file.name)
            document = Document(file=fs.url(filename), content=text)
            document.save()
            process_text_and_index(text, document.id)
        return render(request, 'upload_complete.html', {'documents': Document.objects.all()})
    return render(request, 'upload.html')

#Determines the file type and calls the corresponding function to extract text.
def extract_text(file_path, file_name):
    if file_name.endswith('.pdf'):
        return extract_text_from_pdf(file_path)
    elif file_name.endswith('.docx'):
        return extract_text_from_docx(file_path)
    return "Unsupported file format"

#Extracts text from PDF files
def extract_text_from_pdf(file_path):
    text = ''
    with open(file_path, 'rb') as file:
        pdf_reader = PdfReader(file)
        for page in pdf_reader.pages:
            text += page.extract_text() or ''
    return text

#Extracts text from DOCX files
def extract_text_from_docx(file_path):
    doc = docx.Document(file_path)
    return ' '.join(paragraph.text for paragraph in doc.paragraphs if paragraph.text)

#Processes the extracted text for indexing: tokenizes, stems, and indexes each term.
def process_text_and_index(text, doc_id):
    text = re.sub(r'\W+', ' ', text.lower())
    tokens = word_tokenize(text)
    stemmer = PorterStemmer()
    stemmed_tokens = [stemmer.stem(token) for token in tokens]
    index = {}
    for position, term in enumerate(stemmed_tokens):
        if term not in index:
            index[term] = []
        index[term].append(position)
    for term, positions in index.items():
        pos_index = PositionalIndex(term=term, doc_id=doc_id, positions=','.join(map(str, positions)))
        pos_index.save()

#Cleans and preprocesses the document text for searching: single words are stemmed, phrases are left as is.
def preprocess(document):
    # Clean the document by removing non-word characters (except spaces) and converting to lowercase
    cleaned_document = re.sub(r'[^\w\s]', '', document).lower()
    # Split the document into words
    words = cleaned_document.split()
    # Check if the document contains more than one word
    if len(words) > 1:
        # For multi-word documents, return the cleaned document
        return cleaned_document
    else:
        # For single-word documents, apply stemming
        stemmer = PorterStemmer()
        stemmed_word = stemmer.stem(cleaned_document)
        return stemmed_word

#Handles search queries, processes them, and returns matching documents.
def search(request):
    if request.method == 'POST':
        query = request.POST.get('query', '').strip()
        if ',' in query:
            search_terms = [term.strip() for term in query.split(',')]  # Split the query by commas and strip whitespace
            results = []
            for term in search_terms:
                preprocessed_term = preprocess(term)
                documents = Document.objects.filter(content__icontains=preprocessed_term)
                for doc in documents:
                    starts = [m.start() for m in re.finditer(re.escape(preprocessed_term), doc.content.lower())]
                    results.append({
                        'document_name': doc.file.name,
                        'doc_id': doc.id,
                        'term': preprocessed_term
                    })
        else:
            preprocessed_query = preprocess(query)
            results = []
            documents = Document.objects.filter(content__icontains=preprocessed_query)
            for doc in documents:
                starts = [m.start() for m in re.finditer(re.escape(preprocessed_query), doc.content.lower())]
                results.append({
                    'document_name': doc.file.name,
                    'doc_id': doc.id,
                    'term': preprocessed_query
                })

        return render(request, 'search_results.html', {'results': results, 'query': query})
    return render(request, 'search.html')

#Retrieves and displays a document with highlighted search terms.
def view_document(request, doc_id, term):
    document = Document.objects.get(id=doc_id)
    text = escape(document.content)  # Escape HTML in the original text to prevent XSS
    # Use a regular expression to find all occurrences of the term and apply highlighting
    term_regex = re.compile(re.escape(term), re.IGNORECASE)
    highlighted_text = term_regex.sub(lambda match: f'<span style="background-color: yellow;">{match.group(0)}</span>', text)
    return render(request, 'view_document.html', {'document': document, 'text': mark_safe(highlighted_text)})